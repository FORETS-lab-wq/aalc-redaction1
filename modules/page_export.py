"""
page_export.py
Export du rapport complet en Word (.docx) via python-docx.
Structure : page de garde → profils → analyses individuelles → longitudinales
            → transversale → synthèse H2 → annexes
"""

import io
import datetime
import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from modules.llm import stats_segments

# ─────────────────────────────────────────────────────────────────────────────
# HELPERS DOCX
# ─────────────────────────────────────────────────────────────────────────────

def _rgb(h: str) -> RGBColor:
    h = h.lstrip("#")
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

def _marges(doc: Document):
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(3)
        s.right_margin = Cm(2.5)

def _titre(doc: Document, texte: str, niveau: int = 1):
    p = doc.add_heading(texte, level=niveau)
    if p.runs:
        p.runs[0].font.color.rgb = _rgb("#1B3A5C")
    return p

def _corps(doc: Document, texte: str, italic=False, taille=11, couleur=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(texte)
    run.font.size = Pt(taille)
    run.font.italic = italic
    if couleur:
        run.font.color.rgb = _rgb(couleur)
    return p

def _separation(doc: Document):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E7D9A")
    pBdr.append(bottom)
    pPr.append(pBdr)

def _page_garde(doc: Document, nb_ens: int, date_str: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AALC-Rédaction")
    r.font.size = Pt(24); r.font.bold = True; r.font.color.rgb = _rgb("#1B3A5C")
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Analyse des entretiens semi-directifs\nFormation des enseignants de la conduite automobile — TP ECSR")
    r2.font.size = Pt(14); r2.font.color.rgb = _rgb("#2E7D9A")
    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(
        "Hypothèse 2 : Le dispositif suivi à distance expose les apprenants-enseignants\n"
        "à une rupture de configuration des perspectives d'enseignement."
    )
    r3.font.size = Pt(12); r3.font.italic = True; r3.font.color.rgb = _rgb("#C0392B")
    doc.add_paragraph()
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(f"Rapport complet — N = {nb_ens} enseignants")
    r4.font.size = Pt(16); r4.font.bold = True
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p5.add_run(date_str).font.size = Pt(10)
    doc.add_page_break()

def _tableau_profil(doc: Document, id_e: str, p: dict, stats: dict):
    donnees = [
        ("Identifiant", id_e),
        ("Groupe de formation", p.get("groupe","—")),
        ("Ancienneté professionnelle", p.get("anciennete","—")),
        ("Type de formation", p.get("type_formation","TP ECSR")),
        ("Niveau numérique", p.get("niveau_num","—")),
        ("Phase de l'entretien", p.get("phase","—")),
        ("Organisme", p.get("organisme","—")),
        ("Genre", p.get("genre","—")),
        ("Tranche d'âge", p.get("age_approx","—")),
        ("Segments codés", f"{stats['n']} total — {stats['nR']} ruptures ({stats['pct_r']}%) — {stats['nC']} continuités"),
    ]
    table = doc.add_table(rows=len(donnees), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(donnees):
        row = table.rows[i]
        row.cells[0].text = k
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        # fond bleu léger cellule clé
        tcPr = row.cells[0]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),"D6EAF8")
        tcPr.append(shd)
        row.cells[1].text = str(v)
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()

def _ajouter_texte_redige(doc: Document, texte: str):
    """Découpe le texte rédigé en paragraphes et les ajoute au document."""
    if not texte:
        _corps(doc, "[Texte non généré]", italic=True, couleur="#888888")
        return
    for ligne in texte.split("\n"):
        ligne = ligne.strip()
        if not ligne:
            continue
        if ligne.startswith("### "):
            _titre(doc, ligne[4:], niveau=3)
        elif ligne.startswith("## "):
            _titre(doc, ligne[3:], niveau=2)
        elif ligne.startswith("# "):
            _titre(doc, ligne[2:], niveau=1)
        else:
            p = doc.add_paragraph(ligne)
            p.paragraph_format.space_after = Pt(4)
            for run in p.runs:
                run.font.size = Pt(11)

# ─────────────────────────────────────────────────────────────────────────────
# GÉNÉRATION DU RAPPORT
# ─────────────────────────────────────────────────────────────────────────────

def generer_rapport_word(ens_dict: dict, textes: dict) -> bytes:
    doc = Document()
    _marges(doc)
    date_str = datetime.date.today().strftime("Généré le %d/%m/%Y")
    ids = sorted(ens_dict.keys())

    # ── Page de garde ─────────────────────────────────────────────────────────
    _page_garde(doc, len(ids), date_str)

    # ── Sommaire manuel ───────────────────────────────────────────────────────
    _titre(doc, "Sommaire", 1)
    sommaire = [
        "1. Fiches profils individuels",
        "2. Analyses individuelles (× 10)",
        "3. Analyses longitudinales (× 10)",
        "4. Analyse transversale — Hybride vs Traditionnel",
        "5. Synthèse globale et réponse à l'Hypothèse 2",
    ]
    for item in sommaire:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_page_break()

    # ── PARTIE 1 : Fiches profils ─────────────────────────────────────────────
    _titre(doc, "Partie 1 — Fiches profils individuels", 1)
    _corps(
        doc,
        "Cette partie présente les caractéristiques socioprofessionnelles de chacun des dix enseignants "
        "participants à l'étude, ainsi que leur positionnement initial vis-à-vis du dispositif hybride.",
        italic=True, couleur="#444444"
    )
    doc.add_paragraph()

    for id_e in ids:
        d = ens_dict[id_e]
        p = d.get("profil", {})
        segs = d.get("segments", [])
        stats = stats_segments(segs)

        _titre(doc, f"Fiche profil — {id_e}", 2)
        _tableau_profil(doc, id_e, p, stats)

        cle_profil = f"profil_{id_e}"
        if cle_profil in textes:
            _ajouter_texte_redige(doc, textes[cle_profil])
        _separation(doc)
        doc.add_paragraph()

    doc.add_page_break()

    # ── PARTIE 2 : Analyses individuelles ────────────────────────────────────
    _titre(doc, "Partie 2 — Analyses individuelles", 1)
    _corps(
        doc,
        "Cette partie présente l'analyse détaillée de chaque enseignant, structurée selon les cinq nœuds "
        "du référentiel de codage (Contexte initial, Vécu du dispositif, Alignement théorique, "
        "Transformation des pratiques, Bilan réflexif).",
        italic=True, couleur="#444444"
    )
    doc.add_paragraph()

    for id_e in ids:
        _titre(doc, f"Analyse individuelle — {id_e}", 2)
        cle = f"individuel_{id_e}"
        if cle in textes:
            _ajouter_texte_redige(doc, textes[cle])
        else:
            _corps(doc, "[Analyse non générée — lancez la génération dans le module Traitement individuel]",
                   italic=True, couleur="#888888")
        doc.add_page_break()

    # ── PARTIE 3 : Analyses longitudinales ───────────────────────────────────
    _titre(doc, "Partie 3 — Analyses longitudinales", 1)
    _corps(
        doc,
        "Cette partie reconstitue, pour chaque enseignant, l'évolution de son discours depuis "
        "ses dispositions initiales jusqu'à son bilan réflexif final.",
        italic=True, couleur="#444444"
    )
    doc.add_paragraph()

    for id_e in ids:
        _titre(doc, f"Analyse longitudinale — {id_e}", 2)
        cle = f"longitudinal_{id_e}"
        if cle in textes:
            _ajouter_texte_redige(doc, textes[cle])
        else:
            _corps(doc, "[Analyse longitudinale non générée]", italic=True, couleur="#888888")
        doc.add_page_break()

    # ── PARTIE 4 : Analyse transversale ──────────────────────────────────────
    _titre(doc, "Partie 4 — Analyse transversale : Hybride vs Traditionnel", 1)
    if "transversal" in textes:
        _ajouter_texte_redige(doc, textes["transversal"])
    else:
        _corps(doc, "[Analyse transversale non générée]", italic=True, couleur="#888888")
    doc.add_page_break()

    # ── PARTIE 5 : Synthèse & H2 ─────────────────────────────────────────────
    _titre(doc, "Partie 5 — Synthèse globale et réponse à l'Hypothèse 2", 1)
    if "synthese" in textes:
        _ajouter_texte_redige(doc, textes["synthese"])
    else:
        _corps(doc, "[Synthèse non générée]", italic=True, couleur="#888888")
    doc.add_page_break()

    # ── Annexe : segments codés ───────────────────────────────────────────────
    _titre(doc, "Annexe — Corpus de segments codés", 1)
    for id_e in ids:
        segs = ens_dict[id_e].get("segments", [])
        if not segs:
            continue
        _titre(doc, id_e, 2)
        PERES = ["CONTEXTE_INITIAL","VECU_DISPOSITIF","ALIGNEMENT_THEORIQUE","TRANSFORMATION_PRATIQUE","BILAN_REFLEXIF"]
        for pere in PERES:
            sub = [s for s in segs if s.get("code_pere")==pere]
            if not sub:
                continue
            _titre(doc, pere, 3)
            for s in sub:
                pol = s.get("polarite","AMBIGU")
                pol_emoji = {"RUPTURE":"🔴","CONTINUITE":"🟢","AMBIGU":"⚪"}.get(pol,"⚪")
                _corps(doc, f"[{s.get('code_fils','?')}] {pol_emoji} {pol}", taille=10)
                p_ext = doc.add_paragraph()
                p_ext.paragraph_format.left_indent = Cm(1)
                r = p_ext.add_run(f"« {s.get('extrait','')} »")
                r.font.italic = True; r.font.size = Pt(10)
                if s.get("justification"):
                    p_just = doc.add_paragraph()
                    p_just.paragraph_format.left_indent = Cm(1)
                    rj = p_just.add_run(f"↳ {s['justification']}")
                    rj.font.size = Pt(9); rj.font.color.rgb = _rgb("#666666")
                doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# PAGE STREAMLIT
# ─────────────────────────────────────────────────────────────────────────────

def render():
    st.markdown("# 📄 Export Word (.docx)")
    st.markdown(
        "Génère le rapport complet de la thèse en un seul fichier Word structuré, "
        "prêt à être intégré dans votre manuscrit."
    )

    data     = st.session_state.get("data_aalc", {})
    ens_dict = data.get("enseignants", {})
    textes   = st.session_state.get("textes_rediges", {})

    if not ens_dict:
        st.warning("⚠️ Aucune donnée chargée.")
        return

    ids = sorted(ens_dict.keys())

    # ── Bilan des textes disponibles ──────────────────────────────────────────
    st.markdown("### Bilan des textes rédigés")
    lignes = []
    for id_e in ids:
        lignes.append({
            "Enseignant": id_e,
            "Groupe": ens_dict[id_e].get("profil",{}).get("groupe","—"),
            "Profil": "✅" if f"profil_{id_e}" in textes else "⬜",
            "Analyse individuelle": "✅" if f"individuel_{id_e}" in textes else "⬜",
            "Longitudinale": "✅" if f"longitudinal_{id_e}" in textes else "⬜",
        })
    import pandas as pd
    df = pd.DataFrame(lignes).set_index("Enseignant")
    st.dataframe(df, use_container_width=True)

    col1, col2, col3 = st.columns(3)
    col1.metric("Analyse transversale", "✅" if "transversal" in textes else "⬜")
    col2.metric("Synthèse H2", "✅" if "synthese" in textes else "⬜")
    nb_total = len(ids)*3 + 2
    nb_ok = sum([
        sum(1 for id_e in ids if f"profil_{id_e}" in textes),
        sum(1 for id_e in ids if f"individuel_{id_e}" in textes),
        sum(1 for id_e in ids if f"longitudinal_{id_e}" in textes),
        1 if "transversal" in textes else 0,
        1 if "synthese" in textes else 0,
    ])
    col3.metric("Complétion", f"{nb_ok}/{nb_total} textes")

    st.markdown("---")

    # ── Export ────────────────────────────────────────────────────────────────
    st.markdown("### Générer le rapport Word")
    st.markdown(
        "Le rapport inclut tous les textes rédigés disponibles. "
        "Les sections non générées apparaîtront avec une mention '[non généré]'."
    )

    if nb_ok < nb_total:
        st.warning(
            f"⚠️ {nb_total - nb_ok} texte(s) non encore généré(s). "
            "Vous pouvez exporter maintenant avec les textes disponibles, "
            "ou compléter les modules manquants avant d'exporter."
        )

    col_a, col_b = st.columns([1, 3])
    with col_a:
        if st.button("📄 Générer le .docx", type="primary"):
            with st.spinner("Génération du rapport Word… (30-60s selon le volume)"):
                try:
                    octets = generer_rapport_word(ens_dict, textes)
                    st.session_state["rapport_bytes"] = octets
                    st.success("✅ Rapport généré !")
                except Exception as e:
                    st.error(f"Erreur : {e}")

    if "rapport_bytes" in st.session_state:
        nom = f"rapport_these_TPECSR_{datetime.date.today()}.docx"
        st.download_button(
            label="⬇️ Télécharger le rapport Word",
            data=st.session_state["rapport_bytes"],
            file_name=nom,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
        )
        st.info(f"📄 Fichier : `{nom}`")

    st.markdown("---")
    st.markdown("### Sauvegarde JSON de session")
    st.markdown("Exportez toutes vos données pour les recharger lors d'une prochaine session.")
    import json
    sauvegarde = json.dumps({
        "data_aalc": st.session_state.get("data_aalc", {}),
        "textes_rediges": textes,
    }, ensure_ascii=False, indent=2)
    st.download_button(
        label="💾 Sauvegarder la session (.json)",
        data=sauvegarde.encode("utf-8"),
        file_name=f"session_aalc_{datetime.date.today()}.json",
        mime="application/json",
    )
